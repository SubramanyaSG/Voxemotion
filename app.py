"""
app.py  –  Flask backend for VoxEmotion TTS (v3)
Fixes:
  - Tacotron2 "Reached max decoder steps": text split into sentences,
    each sentence synthesized independently then concatenated
  - Full text coverage: retrieval fallback maps one audio clip per sentence
  - Fast transforms: numpy interp (no scipy overhead)
  - Audio served with correct MIME type + no-cache headers
  - Duration returned in JSON for instant frontend display
"""

import os, sys, json, re, uuid, traceback
import numpy as np
import soundfile as sf
import scipy.interpolate as interp
from pathlib import Path
from flask import Flask, request, jsonify, render_template, Response
from flask_cors import CORS
import warnings
warnings.filterwarnings('ignore')

# ── soxr / numpy resampler (no pkg_resources) ────────────────────────────────
try:
    import soxr
    def _resample(data, sr_in, sr_out):
        return soxr.resample(data.astype(np.float32), sr_in, sr_out, quality='HQ')
except ImportError:
    def _resample(data, sr_in, sr_out):
        if sr_in == sr_out: return data
        n = int(len(data) * sr_out / sr_in)
        return np.interp(np.linspace(0, len(data)-1, n),
                         np.arange(len(data)), data).astype(np.float32)

# ── Runtime config ────────────────────────────────────────────────────────────
_CFG_PATH = os.path.join(os.path.dirname(__file__), 'runtime_config.json')
_CFG = {}
if os.path.exists(_CFG_PATH):
    with open(_CFG_PATH) as f:
        _CFG = json.load(f)

DATASET_ROOT = os.environ.get('DATASET_ROOT', _CFG.get('DATASET_ROOT', ''))
OUTPUT_DIR   = os.environ.get('OUTPUT_DIR',   _CFG.get('OUTPUT_DIR',   'outputs'))
MODEL_DIR    = os.environ.get('MODEL_DIR',    _CFG.get('MODEL_DIR',    'models'))
SAMPLE_RATE  = int(_CFG.get('SAMPLE_RATE', 22050))
EMOTIONS     = _CFG.get('EMOTIONS', ['angry', 'happy', 'neutral', 'sad', 'surprise'])

os.makedirs(OUTPUT_DIR, exist_ok=True)
os.makedirs(MODEL_DIR,  exist_ok=True)

# ── Tacotron2 max chars per chunk (keep well below decoder limit) ─────────────
T2_MAX_CHARS   = 150   # ~one sentence; Tacotron2 is stable up to ~180 chars
SILENCE_MS     = 180   # ms of silence to insert between sentences
SILENCE_SAMPLES = int(SAMPLE_RATE * SILENCE_MS / 1000)

# ── Audio I/O ─────────────────────────────────────────────────────────────────
def load_audio(path: str, target_sr: int = SAMPLE_RATE,
               max_seconds: float = None) -> tuple:
    data, sr = sf.read(path, dtype='float32', always_2d=False)
    if data.ndim > 1:
        data = data.mean(axis=1)
    if max_seconds is not None:
        data = data[:int(sr * max_seconds)]
    if sr != target_sr:
        data = _resample(data, sr, target_sr)
    return data.astype(np.float32), target_sr

def save_audio(path: str, audio: np.ndarray, sr: int = SAMPLE_RATE):
    sf.write(path, np.clip(audio, -1.0, 1.0).astype(np.float32), sr)

# ── Fast emotion transforms (numpy only) ─────────────────────────────────────
def pitch_shift(audio: np.ndarray, sr: int, n_steps: float) -> np.ndarray:
    if n_steps == 0: return audio
    rate   = 2.0 ** (n_steps / 12.0)
    target = max(1, int(len(audio) / rate))
    shifted = np.interp(np.linspace(0, len(audio)-1, target),
                        np.arange(len(audio)), audio)
    return np.interp(np.linspace(0, len(shifted)-1, len(audio)),
                     np.arange(len(shifted)), shifted).astype(np.float32)

def time_stretch(audio: np.ndarray, rate: float) -> np.ndarray:
    if rate == 1.0: return audio
    target = max(1, int(len(audio) / rate))
    return np.interp(np.linspace(0, len(audio)-1, target),
                     np.arange(len(audio)), audio).astype(np.float32)

def apply_emotion(audio: np.ndarray, sr: int, emotion: str) -> np.ndarray:
    ps, es, ts = EMOTION_PARAMS.get(emotion, (0.0, 1.0, 1.0))
    if ps != 0:   audio = pitch_shift(audio, sr, ps)
    audio = audio * es
    if ts != 1.0: audio = time_stretch(audio, ts)
    return np.clip(audio, -1.0, 1.0).astype(np.float32)

# ── Text utilities ────────────────────────────────────────────────────────────
try:
    import inflect
    from unidecode import unidecode
    _inf = inflect.engine()
    def normalize_text(text: str) -> str:
        text = unidecode(str(text))
        text = re.sub(r'\d+', lambda m: _inf.number_to_words(m.group()), text)
        text = re.sub(r"[^a-zA-Z0-9\s.,!?'\-]", '', text)
        return re.sub(r'\s+', ' ', text).strip()
except ImportError:
    def normalize_text(text: str) -> str:
        text = re.sub(r"[^a-zA-Z0-9\s.,!?'\-]", '', str(text))
        return re.sub(r'\s+', ' ', text).strip()


def split_into_chunks(text: str, max_chars: int = T2_MAX_CHARS) -> list:
    """
    Split text into sentence-level chunks each under max_chars.
    Sentences are split on  .  !  ?  then merged greedily up to max_chars.
    """
    # Split on sentence-ending punctuation, keeping the delimiter
    raw = re.split(r'(?<=[.!?])\s+', text.strip())
    chunks, buf = [], ''
    for sent in raw:
        sent = sent.strip()
        if not sent:
            continue
        if len(buf) + len(sent) + 1 <= max_chars:
            buf = (buf + ' ' + sent).strip() if buf else sent
        else:
            if buf:
                chunks.append(buf)
            # If a single sentence is still too long, split on commas
            if len(sent) > max_chars:
                parts = re.split(r'(?<=,)\s+', sent)
                sub_buf = ''
                for p in parts:
                    if len(sub_buf) + len(p) + 1 <= max_chars:
                        sub_buf = (sub_buf + ' ' + p).strip() if sub_buf else p
                    else:
                        if sub_buf: chunks.append(sub_buf)
                        # Hard split as last resort
                        while len(p) > max_chars:
                            chunks.append(p[:max_chars])
                            p = p[max_chars:]
                        sub_buf = p
                if sub_buf: chunks.append(sub_buf)
            else:
                buf = sent
    if buf:
        chunks.append(buf)
    return [c for c in chunks if c.strip()]

# ── File text extraction ───────────────────────────────────────────────────────
def extract_text_from_pdf(path: str) -> str:
    try:
        from pdfminer.high_level import extract_text
        return extract_text(path)
    except Exception:
        try:
            import PyPDF2
            with open(path, 'rb') as f:
                r = PyPDF2.PdfReader(f)
                return '\n'.join(p.extract_text() or '' for p in r.pages)
        except Exception as e:
            return f'[PDF error: {e}]'

def extract_text_from_docx(path: str) -> str:
    try:
        from docx import Document
        return '\n'.join(p.text for p in Document(path).paragraphs)
    except Exception as e:
        return f'[DOCX error: {e}]'

def extract_text_from_txt(path: str) -> str:
    try:
        with open(path, 'r', encoding='utf-8', errors='replace') as f:
            return f.read()
    except Exception as e:
        return f'[TXT error: {e}]'

# ── Dataset loader ─────────────────────────────────────────────────────────────
import pandas as pd
_df_cache = None

def get_df():
    global _df_cache
    if _df_cache is not None: return _df_cache
    csv = os.path.join(OUTPUT_DIR, 'dataset_metadata.csv')
    if os.path.exists(csv):
        _df_cache = pd.read_csv(csv)
        return _df_cache
    records = []
    if DATASET_ROOT and os.path.isdir(DATASET_ROOT):
        for spk in sorted(Path(DATASET_ROOT).iterdir()):
            if not spk.is_dir(): continue
            for emo in sorted(d for d in spk.iterdir() if d.is_dir()):
                for wav in sorted(emo.glob('*.wav')):
                    try:
                        info = sf.info(str(wav))
                        records.append({'speaker': spk.name, 'emotion': emo.name.lower(),
                                        'file': str(wav), 'filename': wav.name,
                                        'text': '', 'duration': info.duration, 'readable': True})
                    except Exception: pass
    _df_cache = pd.DataFrame(records)
    return _df_cache

# ── Emotion params ─────────────────────────────────────────────────────────────
EMOTION_PARAMS = {
    'angry'   : ( 2.0, 1.30, 1.05),
    'happy'   : ( 3.5, 1.20, 1.10),
    'neutral' : ( 0.0, 1.00, 1.00),
    'sad'     : (-3.0, 0.75, 0.88),
    'surprise': ( 4.5, 1.15, 1.15),
}

# ── Tacotron2 model cache (loaded once, reused across requests) ───────────────
_t2_model  = None
_wg_model  = None
_t2_utils  = None
_t2_failed = False

def _load_tacotron2():
    global _t2_model, _wg_model, _t2_utils, _t2_failed
    if _t2_failed or _t2_model is not None:
        return _t2_model is not None
    try:
        import torch
        dev = torch.device('cuda' if torch.cuda.is_available() else 'cpu')
        print('⏳  Loading Tacotron2 (once) …')
        _t2_model = torch.hub.load('NVIDIA/DeepLearningExamples:torchhub',
                                    'nvidia_tacotron2', model_math='fp32',
                                    verbose=False).to(dev).eval()
        _wg_model = torch.hub.load('NVIDIA/DeepLearningExamples:torchhub',
                                    'nvidia_waveglow', model_math='fp32',
                                    verbose=False).to(dev).eval()
        _t2_utils = torch.hub.load('NVIDIA/DeepLearningExamples:torchhub',
                                    'nvidia_tts_utils', verbose=False)
        for m in _wg_model.modules():
            if hasattr(m, 'weight_v'):
                import torch.nn.utils as nnu
                nnu.remove_weight_norm(m)
        print('✅  Tacotron2 loaded.')
        return True
    except Exception as e:
        print(f'⚠  Tacotron2 unavailable: {e}')
        _t2_failed = True
        return False


def _t2_synth_chunk(chunk: str) -> np.ndarray:
    """Synthesize one short text chunk via Tacotron2. Returns float32 array at 22050 Hz."""
    import torch
    dev = torch.device('cuda' if torch.cuda.is_available() else 'cpu')
    seqs, lengths = _t2_utils.prepare_input_sequence([chunk])
    with torch.no_grad():
        mel, _, _ = _t2_model.infer(seqs.to(dev), lengths.to(dev))
        audio     = _wg_model.infer(mel)
    return audio[0].cpu().numpy().astype(np.float32)


def _retrieval_chunk(emotion: str, duration_hint: float = 3.0) -> np.ndarray:
    """
    Pick a random ESD sample matching the emotion and return audio
    trimmed to duration_hint seconds (so retrieval matches text length roughly).
    """
    df = get_df()
    readable_col = 'readable' if 'readable' in df.columns else None
    if readable_col:
        sub = df[(df[readable_col] == True) & (df['emotion'] == emotion)]
        if sub.empty:
            sub = df[df[readable_col] == True]
    else:
        sub = df[df['emotion'] == emotion] if not df.empty else pd.DataFrame()
        if sub.empty: sub = df

    if sub.empty:
        return np.zeros(int(SAMPLE_RATE * duration_hint), dtype=np.float32)

    try:
        audio, _ = load_audio(sub.sample(1).iloc[0]['file'],
                               SAMPLE_RATE, max_seconds=duration_hint)
        # Pad if shorter than requested duration
        needed = int(SAMPLE_RATE * duration_hint)
        if len(audio) < needed:
            audio = np.pad(audio, (0, needed - len(audio)))
        return audio
    except Exception:
        return np.zeros(int(SAMPLE_RATE * duration_hint), dtype=np.float32)


def synthesize_audio(text: str, emotion: str) -> dict:
    """
    Main synthesis pipeline:
      1. Normalize and split text into short chunks (≤150 chars each)
      2. For each chunk: try Tacotron2 → fallback to retrieval
      3. Concatenate all chunks with silence gaps
      4. Apply emotion transform once on the full audio
      5. Save and return metadata
    """
    emotion = emotion.lower() if emotion.lower() in EMOTION_PARAMS else 'neutral'
    text    = normalize_text(text)[:10_000_000]
    chunks  = split_into_chunks(text, max_chars=T2_MAX_CHARS)

    if not chunks:
        chunks = [text[:T2_MAX_CHARS]]

    use_t2 = _load_tacotron2()

    # Estimate words-per-second for retrieval duration hint (~2.5 words/sec)
    WPS = 2.5
    silence = np.zeros(SILENCE_SAMPLES, dtype=np.float32)
    parts   = []

    print(f'  Synthesizing {len(chunks)} chunk(s) for emotion={emotion}')

    for i, chunk in enumerate(chunks):
        chunk = chunk.strip()
        if not chunk:
            continue
        print(f'  Chunk {i+1}/{len(chunks)}: "{chunk[:60]}{"…" if len(chunk)>60 else ""}"')

        seg = None
        if use_t2:
            try:
                seg = _t2_synth_chunk(chunk)
            except Exception as e:
                print(f'    Tacotron2 chunk failed: {e}')
                seg = None

        if seg is None:
            # Retrieval: estimate how long this sentence should be
            word_count   = len(chunk.split())
            dur_hint     = max(2.0, word_count / WPS)
            seg = _retrieval_chunk(emotion, duration_hint=dur_hint)

        parts.append(seg)
        if i < len(chunks) - 1:
            parts.append(silence)   # gap between sentences

    if not parts:
        audio = np.zeros(SAMPLE_RATE * 3, dtype=np.float32)
        sr    = SAMPLE_RATE
    else:
        audio = np.concatenate(parts).astype(np.float32)
        sr    = 22050 if use_t2 else SAMPLE_RATE

    # Apply emotion transform on the final concatenated audio
    audio = apply_emotion(audio, sr, emotion)

    fname    = f'tts_{emotion}_{uuid.uuid4().hex[:8]}.wav'
    out_path = os.path.join(OUTPUT_DIR, fname)
    save_audio(out_path, audio, sr)

    duration = round(len(audio) / sr, 2)
    print(f'  ✅ Done — {duration:.1f}s saved to {fname}')
    return {'file': fname, 'text': text, 'emotion': emotion, 'duration': duration}


# ── Flask app ──────────────────────────────────────────────────────────────────
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
app = Flask(__name__,
            template_folder=os.path.join(BASE_DIR, 'templates'),
            static_folder=os.path.join(BASE_DIR, 'static'))
<<<<<<< Updated upstream
CORS(app)
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024
=======

app.secret_key = SECRET_KEY
app.config.update(
    SESSION_COOKIE_HTTPONLY   = True,
    SESSION_COOKIE_SAMESITE   = 'Lax',
    SESSION_COOKIE_SECURE     = False,   # Set True in production with HTTPS
    PERMANENT_SESSION_LIFETIME= timedelta(days=7),
    MAX_CONTENT_LENGTH        = 50 * 1024 * 1024,   # 50 MB
)
CORS(app, supports_credentials=True)


# ── Security headers on every response ───────────────────────────────────────
@app.after_request
def set_security_headers(response):
    response.headers['X-Content-Type-Options']  = 'nosniff'
    response.headers['X-Frame-Options']         = 'DENY'
    response.headers['X-XSS-Protection']        = '1; mode=block'
    response.headers['Referrer-Policy']         = 'strict-origin-when-cross-origin'
    response.headers['Permissions-Policy']      = 'microphone=(), camera=()'
    response.headers['Content-Security-Policy'] = (
        "default-src 'self'; "
        "script-src 'self' 'unsafe-inline' https://fonts.googleapis.com; "
        "style-src 'self' 'unsafe-inline' https://fonts.googleapis.com "
        "https://fonts.gstatic.com; "
        "font-src 'self' https://fonts.gstatic.com; "
        "connect-src 'self'; "
        "img-src 'self' data:; "
        "media-src 'self' blob:;"
    )
    return response


# ── Login required decorator ──────────────────────────────────────────────────
def login_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if 'user_email' not in session:
            flash('Please log in to continue.', 'info')
            return redirect(url_for('login_page'))
        return f(*args, **kwargs)
    return decorated


# ════════════════════════════════════════════════════════════════════════════
# AUTH ROUTES
# ════════════════════════════════════════════════════════════════════════════

@app.route('/login', methods=['GET', 'POST'])
def login_page():
    """Login page — also contains register and forgot-password panels."""
    if 'user_email' in session:
        return redirect(url_for('index'))

    if request.method == 'POST':
        # Validate against token minted on the previous page render.
        if not validate_csrf(request.form.get('csrf_token', '')):
            flash('Invalid request. Please try again.', 'error')
            return redirect(url_for('login_page'))

        ip = request.remote_addr
        if not is_rate_allowed(ip, limit=10, window=60):
            flash('Too many login attempts. Please wait 1 minute.', 'error')
            return render_template('login.html', csrf_token=generate_csrf())

        email    = request.form.get('email', '').strip().lower()
        password = request.form.get('password', '')

        if not valid_email(email) or not password:
            flash('Please enter a valid email and password.', 'error')
            return render_template('login.html', csrf_token=generate_csrf())

        user = verify_user(email, password)
        if not user:
            flash('Incorrect email or password.', 'error')
            return render_template('login.html', csrf_token=generate_csrf())

        session.permanent = bool(request.form.get('remember'))
        session['user_email']    = user['email']
        session['user_fullname'] = user.get('fullname', '')
        session['user_role']     = user.get('role', 'user')
        session.pop('csrf_token', None)
        return redirect(url_for('index'))

    return render_template('login.html', csrf_token=generate_csrf())


@app.route('/register', methods=['POST'])
def register():
    """Handle new user registration."""
    if not validate_csrf(request.form.get('csrf_token', '')):
        flash('Invalid request. Please try again.', 'error')
        return redirect(url_for('login_page'))

    ip   = request.remote_addr
    wait = get_reg_cooldown(ip)
    if wait:
        flash(f'Please wait {wait} more seconds before registering again.', 'warning')
        return render_template('login.html', csrf_token=generate_csrf())

    if not is_rate_allowed(ip, limit=3, window=3600):
        flash('Registration limit reached. Please try again later.', 'error')
        return render_template('login.html', csrf_token=generate_csrf())

    fullname = request.form.get('fullname', '').strip()
    dob      = request.form.get('dob', '').strip()
    email    = request.form.get('email', '').strip().lower()
    password = request.form.get('password', '')
    confirm  = request.form.get('confirm_password', '')

    errors = []
    if not valid_fullname(fullname):
        errors.append("Name can include letters, spaces, hyphens, and apostrophes only.")
    if not valid_dob(dob):
        errors.append('Invalid date of birth. Use DD/MM/YYYY.')
    if not valid_email(email):
        errors.append('Invalid email address.')
    if not valid_password(password):
        errors.append('Password must be 8+ characters with 1 uppercase letter and 1 number.')
    if password != confirm:
        errors.append('Passwords do not match.')
    if not request.form.get('agree_terms'):
        errors.append('You must accept the Terms & Conditions.')

    if errors:
        for err in errors:
            flash(err, 'error')
        return render_template('login.html', csrf_token=generate_csrf())

    if user_exists(email):
        flash('An account with this email already exists.', 'error')
        return render_template('login.html', csrf_token=generate_csrf())

    if not create_user(email, password, fullname, dob):
        flash('Registration failed. Please try again.', 'error')
        return render_template('login.html', csrf_token=generate_csrf())

    set_reg_cooldown(ip)
    flash('Account created successfully! Please log in.', 'success')
    return redirect(url_for('login_page'))


@app.route('/forgot-password', methods=['POST'])
def forgot_password():
    """Send password reset email."""
    if not validate_csrf(request.form.get('csrf_token', '')):
        flash('Invalid request.', 'error')
        return redirect(url_for('login_page'))

    if not is_rate_allowed(request.remote_addr, limit=3, window=300):
        flash('Too many reset requests. Please wait 5 minutes.', 'error')
        return render_template('login.html', csrf_token=generate_csrf())

    email = request.form.get('email', '').strip().lower()
    if not valid_email(email):
        flash('Please enter a valid email address.', 'error')
        return render_template('login.html', csrf_token=generate_csrf())

    # Always show same message to prevent email enumeration
    if user_exists(email):
        token = create_reset_token(email)
        send_password_reset_email(email, token)

    flash('If that email is registered, a reset link has been sent.', 'success')
    return redirect(url_for('login_page'))


@app.route('/reset-password/<token>', methods=['GET', 'POST'])
def reset_password(token):
    """Password reset page (accessed via emailed link)."""
    if request.method == 'GET':
        if not is_reset_token_valid(token):
            flash('This reset link is invalid or has expired. Please request a new one.', 'error')
            return redirect(url_for('login_page'))
        return render_template('reset_password.html', token=token, csrf_token=generate_csrf())

    if not validate_csrf(request.form.get('csrf_token', '')):
        flash('Invalid request.', 'error')
        return redirect(url_for('login_page'))

    password = request.form.get('password', '')
    confirm  = request.form.get('confirm_password', '')

    if not valid_password(password):
        flash('Password must be 8+ characters with 1 uppercase letter and 1 number.', 'error')
        return render_template('reset_password.html', token=token, csrf_token=generate_csrf())

    if password != confirm:
        flash('Passwords do not match.', 'error')
        return render_template('reset_password.html', token=token, csrf_token=generate_csrf())

    email = consume_reset_token(token)
    if not email:
        flash('Reset link expired or already used. Please request a new one.', 'error')
        return redirect(url_for('login_page'))

    if update_password(email, password):
        flash('Password updated successfully! Please log in.', 'success')
    else:
        flash('Could not update password. Please try again.', 'error')

    return redirect(url_for('login_page'))


@app.route('/logout')
def logout():
    session.clear()
    flash('You have been signed out.', 'info')
    return redirect(url_for('login_page'))


# ════════════════════════════════════════════════════════════════════════════
# MAIN APP ROUTES  (protected — must be logged in)
# ════════════════════════════════════════════════════════════════════════════
>>>>>>> Stashed changes

@app.route('/')
def index():
    return render_template('index.html', emotions=EMOTIONS)

@app.route('/synthesize', methods=['POST'])
def synthesize():
    try:
        data    = request.get_json(force=True)
        text    = data.get('text', '').strip()
        emotion = data.get('emotion', 'neutral')
        if not text:
            return jsonify({'error': 'No text provided'}), 400
        return jsonify({'success': True, **synthesize_audio(text, emotion)})
    except Exception as e:
        return jsonify({'error': str(e), 'trace': traceback.format_exc()}), 500

@app.route('/upload', methods=['POST'])
def upload():
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file uploaded'}), 400
        f       = request.files['file']
        emotion = request.form.get('emotion', 'neutral')
        ext     = Path(f.filename).suffix.lower()
        if ext not in {'.txt', '.pdf', '.docx'}:
            return jsonify({'error': f'Unsupported type: {ext}'}), 400
        tmp = os.path.join(OUTPUT_DIR, f'upload_{uuid.uuid4().hex}{ext}')
        f.save(tmp)
        text = (extract_text_from_pdf(tmp)  if ext == '.pdf'  else
                extract_text_from_docx(tmp) if ext == '.docx' else
                extract_text_from_txt(tmp))
        os.remove(tmp)
        if not text.strip():
            return jsonify({'error': 'No text extracted from file'}), 400
        return jsonify({'success': True, **synthesize_audio(text, emotion)})
    except Exception as e:
        return jsonify({'error': str(e), 'trace': traceback.format_exc()}), 500

@app.route('/audio/<filename>')
def serve_audio(filename):
    filepath = os.path.join(OUTPUT_DIR, filename)
    if not os.path.exists(filepath):
        return jsonify({'error': 'Audio file not found'}), 404
    with open(filepath, 'rb') as f:
        data = f.read()
    resp = Response(data, mimetype='audio/wav')
    resp.headers['Content-Length']              = len(data)
    resp.headers['Accept-Ranges']               = 'bytes'
    resp.headers['Cache-Control']               = 'no-cache, no-store, must-revalidate'
    resp.headers['Access-Control-Allow-Origin'] = '*'
    return resp

if __name__ == '__main__':
    print('🎙  VoxEmotion Web App  v3 — sentence-chunked synthesis')
    print(f'   Dataset : {DATASET_ROOT}')
    print(f'   Outputs : {OUTPUT_DIR}')
    # Pre-load Tacotron2 at startup so first request is fast
    _load_tacotron2()
    app.run(host='0.0.0.0', port=5000, debug=False, use_reloader=False)

# ─── Runtime config ───────────────────────────────────────────────────────────
_CFG_PATH = os.path.join(os.path.dirname(__file__), 'runtime_config.json')
_CFG = {}
if os.path.exists(_CFG_PATH):
    with open(_CFG_PATH) as f:
        _CFG = json.load(f)

DATASET_ROOT = os.environ.get('DATASET_ROOT', _CFG.get('DATASET_ROOT', ''))
OUTPUT_DIR   = os.environ.get('OUTPUT_DIR',   _CFG.get('OUTPUT_DIR',   'outputs'))
MODEL_DIR    = os.environ.get('MODEL_DIR',    _CFG.get('MODEL_DIR',    'models'))
SAMPLE_RATE  = int(_CFG.get('SAMPLE_RATE', 22050))
EMOTIONS     = _CFG.get('EMOTIONS', ['angry', 'happy', 'neutral', 'sad', 'surprise'])

os.makedirs(OUTPUT_DIR, exist_ok=True)
os.makedirs(MODEL_DIR,  exist_ok=True)

# ─── Audio utilities (pkg_resources-free) ────────────────────────────────────
N_FFT = 1024; HOP = 256; N_MELS = 80
MAX_CLIP_SECONDS = 6   # clip retrieval audio to 6 s for fast transforms

def load_audio(path: str, target_sr: int = SAMPLE_RATE, max_seconds: float = None):
    """Read WAV → mono float32. Clips to max_seconds if given (speeds up transforms)."""
    data, sr = sf.read(path, dtype='float32', always_2d=False)
    if data.ndim > 1:
        data = data.mean(axis=1)
    # Clip early BEFORE resampling – much faster
    if max_seconds is not None:
        max_samples = int(sr * max_seconds)
        data = data[:max_samples]
    if sr != target_sr:
        data = _resample(data, sr, target_sr)
    return data.astype(np.float32), target_sr

def save_audio(path: str, audio: np.ndarray, sr: int = SAMPLE_RATE):
    sf.write(path, np.clip(audio, -1.0, 1.0).astype(np.float32), sr)

def pitch_shift_simple(audio: np.ndarray, sr: int, n_steps: float) -> np.ndarray:
    """Fast pitch shift using numpy (no scipy interp overhead for small arrays)."""
    if n_steps == 0: return audio
    rate   = 2.0 ** (n_steps / 12.0)
    target = max(1, int(len(audio) / rate))
    # Use numpy interp – much faster than scipy.interpolate.interp1d
    shifted = np.interp(np.linspace(0, len(audio)-1, target),
                        np.arange(len(audio)), audio)
    result  = np.interp(np.linspace(0, len(shifted)-1, len(audio)),
                        np.arange(len(shifted)), shifted)
    return result.astype(np.float32)

def time_stretch_simple(audio: np.ndarray, rate: float) -> np.ndarray:
    """Fast time stretch using numpy interp."""
    if rate == 1.0: return audio
    target = max(1, int(len(audio) / rate))
    return np.interp(np.linspace(0, len(audio)-1, target),
                     np.arange(len(audio)), audio).astype(np.float32)

# ─── Text extraction ──────────────────────────────────────────────────────────
def extract_text_from_pdf(path: str) -> str:
    try:
        from pdfminer.high_level import extract_text
        return extract_text(path)
    except Exception:
        try:
            import PyPDF2
            with open(path, 'rb') as f:
                r = PyPDF2.PdfReader(f)
                return '\n'.join(p.extract_text() or '' for p in r.pages)
        except Exception as e:
            return f'[PDF error: {e}]'

def extract_text_from_docx(path: str) -> str:
    try:
        from docx import Document
        return '\n'.join(p.text for p in Document(path).paragraphs)
    except Exception as e:
        return f'[DOCX error: {e}]'

def extract_text_from_txt(path: str) -> str:
    try:
        with open(path, 'r', encoding='utf-8', errors='replace') as f:
            return f.read()
    except Exception as e:
        return f'[TXT error: {e}]'

# ─── Text normalisation ───────────────────────────────────────────────────────
try:
    import inflect
    from Unidecode import unidecode
    _inf = inflect.engine()
    def normalize_text(text: str) -> str:
        text = unidecode(str(text))
        text = re.sub(r'\d+', lambda m: _inf.number_to_words(m.group()), text)
        text = re.sub(r"[^a-zA-Z0-9\s.,!?'\-]", '', text)
        return re.sub(r'\s+', ' ', text).strip()
except ImportError:
    def normalize_text(text: str) -> str:
        text = re.sub(r"[^a-zA-Z0-9\s.,!?'\-]", '', str(text))
        return re.sub(r'\s+', ' ', text).strip()

# ─── Dataset loader ───────────────────────────────────────────────────────────
import pandas as pd
_df_cache = None

def get_df():
    global _df_cache
    if _df_cache is not None: return _df_cache
    csv = os.path.join(OUTPUT_DIR, 'dataset_metadata.csv')
    if os.path.exists(csv):
        _df_cache = pd.read_csv(csv)
        return _df_cache
    # Fallback scan
    records = []
    if DATASET_ROOT and os.path.isdir(DATASET_ROOT):
        for spk in sorted(Path(DATASET_ROOT).iterdir()):
            if not spk.is_dir(): continue
            for emo in sorted(d for d in spk.iterdir() if d.is_dir()):
                for wav in sorted(emo.glob('*.wav')):
                    try:
                        info = sf.info(str(wav))
                        records.append({'speaker':spk.name,'emotion':emo.name.lower(),
                                        'file':str(wav),'filename':wav.name,
                                        'text':'','duration':info.duration,'readable':True})
                    except Exception: pass
    _df_cache = pd.DataFrame(records)
    return _df_cache

# ─── Emotion params ───────────────────────────────────────────────────────────
EMOTION_PARAMS = {
    'angry'   : ( 2.0, 1.30, 1.05),
    'happy'   : ( 3.5, 1.20, 1.10),
    'neutral' : ( 0.0, 1.00, 1.00),
    'sad'     : (-3.0, 0.75, 0.88),
    'surprise': ( 4.5, 1.15, 1.15),
}

def synthesize_audio(text: str, emotion: str) -> dict:
    df      = get_df()
    emotion = emotion.lower() if emotion.lower() in EMOTION_PARAMS else 'neutral'
    text    = normalize_text(text)[:10_000_000]

    audio = None; sr = SAMPLE_RATE

    # Try Tacotron2 (skip if already known to fail via env flag)
    if os.environ.get('TACOTRON2_FAILED') != '1':
        try:
            import torch
            dev = torch.device('cuda' if torch.cuda.is_available() else 'cpu')
            t2  = torch.hub.load('NVIDIA/DeepLearningExamples:torchhub',
                                 'nvidia_tacotron2', model_math='fp32', verbose=False).to(dev).eval()
            wg  = torch.hub.load('NVIDIA/DeepLearningExamples:torchhub',
                                 'nvidia_waveglow',  model_math='fp32', verbose=False).to(dev).eval()
            ut  = torch.hub.load('NVIDIA/DeepLearningExamples:torchhub', 'nvidia_tts_utils', verbose=False)
            for m in wg.modules():
                if hasattr(m, 'weight_v'): torch.nn.utils.remove_weight_norm(m)
            s, l = ut.prepare_input_sequence([text[:500]])
            with torch.no_grad():
                mel, _, _ = t2.infer(s.to(dev), l.to(dev))
                audio     = wg.infer(mel)[0].cpu().numpy().astype(np.float32)
            sr = 22050
        except Exception:
            os.environ['TACOTRON2_FAILED'] = '1'   # skip on next call
            audio = None

    # Fast retrieval fallback
    if audio is None:
        readable_col = 'readable' if 'readable' in df.columns else None
        if readable_col:
            sub = df[(df[readable_col] == True) & (df['emotion'] == emotion)]
            if sub.empty:
                sub = df[df[readable_col] == True]
        else:
            sub = df[df['emotion'] == emotion] if not df.empty else pd.DataFrame()
            if sub.empty: sub = df

        if sub.empty:
            audio = np.zeros(sr * 3, dtype=np.float32)
        else:
            try:
                # ── KEY FIX: load only first 6 seconds for fast transforms ──
                audio, sr = load_audio(sub.sample(1).iloc[0]['file'],
                                       SAMPLE_RATE, max_seconds=MAX_CLIP_SECONDS)
            except Exception:
                audio = np.zeros(SAMPLE_RATE * 3, dtype=np.float32)
                sr    = SAMPLE_RATE

    # Apply emotion transform (fast numpy version)
    ps, es, ts = EMOTION_PARAMS.get(emotion, (0.0, 1.0, 1.0))
    if ps != 0:   audio = pitch_shift_simple(audio, sr, ps)
    audio = audio * es
    if ts != 1.0: audio = time_stretch_simple(audio, ts)
    audio = np.clip(audio, -1.0, 1.0).astype(np.float32)

    fname    = f'tts_{emotion}_{uuid.uuid4().hex[:8]}.wav'
    out_path = os.path.join(OUTPUT_DIR, fname)
    save_audio(out_path, audio, sr)

    duration = round(len(audio) / sr, 2)
    return {'file': fname, 'text': text, 'emotion': emotion, 'duration': duration}

# ─── Flask app ────────────────────────────────────────────────────────────────
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
app = Flask(__name__,
            template_folder=os.path.join(BASE_DIR, 'templates'),
            static_folder=os.path.join(BASE_DIR, 'static'))
CORS(app)
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024   # 50 MB

@app.route('/')
def index():
    return render_template('index.html', emotions=EMOTIONS)

@app.route('/synthesize', methods=['POST'])
def synthesize():
    try:
        data    = request.get_json(force=True)
        text    = data.get('text','').strip()
        emotion = data.get('emotion','neutral')
        if not text: return jsonify({'error':'No text provided'}), 400
        return jsonify({'success':True, **synthesize_audio(text, emotion)})
    except Exception as e:
        return jsonify({'error':str(e), 'trace':traceback.format_exc()}), 500

@app.route('/upload', methods=['POST'])
def upload():
    try:
        if 'file' not in request.files:
            return jsonify({'error':'No file uploaded'}), 400
        f       = request.files['file']
        emotion = request.form.get('emotion','neutral')
        ext     = Path(f.filename).suffix.lower()
        if ext not in {'.txt','.pdf','.docx'}:
            return jsonify({'error':f'Unsupported type: {ext}'}), 400
        tmp = os.path.join(OUTPUT_DIR, f'upload_{uuid.uuid4().hex}{ext}')
        f.save(tmp)
        text = (extract_text_from_pdf(tmp)   if ext=='.pdf'  else
                extract_text_from_docx(tmp)  if ext=='.docx' else
                extract_text_from_txt(tmp))
        os.remove(tmp)
        if not text.strip(): return jsonify({'error':'No text extracted'}), 400
        return jsonify({'success':True, **synthesize_audio(text, emotion)})
    except Exception as e:
        return jsonify({'error':str(e), 'trace':traceback.format_exc()}), 500

@app.route('/audio/<filename>')
def serve_audio(filename):
    # Serve with explicit WAV MIME type and no-cache headers
    filepath = os.path.join(OUTPUT_DIR, filename)
    if not os.path.exists(filepath):
        return jsonify({'error': 'Audio file not found'}), 404
    with open(filepath, 'rb') as f:
        data = f.read()
    resp = Response(data, mimetype='audio/wav')
    resp.headers['Content-Length']       = len(data)
    resp.headers['Accept-Ranges']        = 'bytes'
    resp.headers['Cache-Control']        = 'no-cache, no-store, must-revalidate'
    resp.headers['Access-Control-Allow-Origin'] = '*'
    return resp

if __name__ == '__main__':
    print('🎙  VoxEmotion Web App  (pkg_resources-free)')
    print(f'   Dataset : {DATASET_ROOT}')
    print(f'   Outputs : {OUTPUT_DIR}')
    app.run(host='0.0.0.0', port=5000, debug=False, use_reloader=False)
